"""One-off script that generates templates_docx/handover_template.docx.

Starts from the *actual* "HAND OVER FORM - 2025.docx" and surgically
replaces only the variable values with Jinja placeholders - run by run - so
the original fonts, bold labels, column widths, and the Shinhan logo in the
page header all carry over untouched. Editing in place is safe here because
every value we need to templatize already sits in its own dedicated run in
the source file (verified by inspection); we only ever set `run.text`, never
touch surrounding runs, so nothing else shifts.

Run with: python scripts/build_handover_template.py
"""

import os
import shutil
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

SOURCE_PATH = os.path.join(os.path.dirname(__file__), "..", "HAND OVER FORM - 2025.docx")
OUT_PATH = os.path.join(os.path.dirname(__file__), "..", "templates_docx", "handover_template.docx")


def set_cell_text(cell, text):
    """Used only for the docxtpl for/endfor wrapper rows, whose content is
    entirely discarded at render time - formatting doesn't matter there."""
    cell.text = ""
    cell.paragraphs[0].add_run(text)


def replace_run_text(paragraph, index, new_text):
    paragraph.runs[index].text = new_text


def main():
    shutil.copy(SOURCE_PATH, OUT_PATH)
    doc = Document(OUT_PATH)

    # --- "Today, __/__/2026, here we are:" -> real date placeholders --------
    # Original runs: ' ', 'Today', ',', '.…..', '/……/ ', '202', '6', ', here we are:'
    date_p = doc.paragraphs[2]
    assert "Today" in date_p.text, f"Unexpected paragraph 2: {date_p.text!r}"
    date_p.runs[3].text = " {{ ho_day }}"
    date_p.runs[4].text = "/{{ ho_month }}/"
    date_p.runs[5].text = "{{ ho_year }}"
    date_p.runs[6].text = ""

    # --- Table 0: ICT representative + receiving party ----------------------
    t0 = doc.tables[0]
    replace_run_text(t0.rows[1].cells[1].paragraphs[0], 0, "{{ ict_rep_name }}")
    replace_run_text(t0.rows[1].cells[3].paragraphs[0], 0, "{{ ict_rep_id }}")
    replace_run_text(t0.rows[3].cells[1].paragraphs[0], 0, "{{ receiving_name }}")
    replace_run_text(t0.rows[3].cells[3].paragraphs[0], 0, "{{ receiving_title }}")
    replace_run_text(t0.rows[4].cells[1].paragraphs[0], 0, "{{ receiving_dept }}")
    replace_run_text(t0.rows[4].cells[3].paragraphs[0], 0, "{{ receiving_id }}")

    # --- Table 1: type of hand-over checkboxes -------------------------------
    # ASSIGNMENT/TEMP = giving equipment out (hand over); RETURN = getting it
    # back (receive) - the direction is conveyed by these existing labels
    # rather than a separate field (confirmed with the requester).
    t1 = doc.tables[1]
    replace_run_text(t1.rows[0].cells[1].paragraphs[0], 0, "{{ type_assignment_box }}")
    replace_run_text(t1.rows[0].cells[2].paragraphs[0], 0, "{{ type_temp_box }}")
    due_date_runs = t1.rows[0].cells[2].paragraphs[0].runs
    due_date_runs[-1].text = due_date_runs[-1].text + "{{ temp_due_date }}"
    replace_run_text(t1.rows[1].cells[1].paragraphs[0], 0, "{{ type_return_box }}")
    replace_run_text(t1.rows[1].cells[2].paragraphs[0], 0, "{{ type_other_box }}")
    other_type_runs = t1.rows[1].cells[2].paragraphs[0].runs
    other_type_runs[-1].text = other_type_runs[-1].text + " {{ other_type_text }}"

    # --- Table 2: reason of hand-over checkboxes -----------------------------
    t2 = doc.tables[2]
    replace_run_text(t2.rows[0].cells[1].paragraphs[0], 0, "{{ reason_newcomer_box }}")
    replace_run_text(t2.rows[0].cells[2].paragraphs[0], 0, "{{ reason_rotate_box }}")
    replace_run_text(t2.rows[0].cells[3].paragraphs[0], 0, "{{ reason_maternity_box }}")
    replace_run_text(t2.rows[1].cells[1].paragraphs[0], 0, "{{ reason_resigned_box }}")
    replace_run_text(t2.rows[1].cells[2].paragraphs[0], 0, "{{ reason_other_box }}")
    other_reason_runs = t2.rows[1].cells[2].paragraphs[0].runs
    other_reason_runs[-1].text = other_reason_runs[-1].text + " {{ other_reason_text }}"
    # Note: cells[2] and cells[3] on this row are the same merged cell in the
    # source file (one wide "OTHER" reason box spanning two grid columns),
    # not two independent fields.

    # --- Table 3: equipment list (docxtpl row-loop) --------------------------
    # docxtpl's `{%tr %}` tag replaces the *entire* <w:tr> row it's found in,
    # so the for/endfor markers each need their own dedicated (throwaway) row
    # cloned from the real data row, with that original row left untouched
    # in between so its <w:tr> XML is what actually gets repeated.
    t3 = doc.tables[3]
    data_row = t3.rows[1]
    tr_data = data_row._tr
    tr_for = deepcopy(tr_data)
    tr_endfor = deepcopy(tr_data)
    tr_data.addprevious(tr_for)
    tr_data.addnext(tr_endfor)

    t3 = doc.tables[3]  # re-read after structural edit
    for_row, data_row, endfor_row = t3.rows[1], t3.rows[2], t3.rows[3]

    set_cell_text(for_row.cells[0], "{%tr for item in assets %}")
    for c in for_row.cells[1:]:
        set_cell_text(c, "")
    set_cell_text(endfor_row.cells[0], "{%tr endfor %}")
    for c in endfor_row.cells[1:]:
        set_cell_text(c, "")

    replace_run_text(data_row.cells[0].paragraphs[0], 0, "{{ item.no }}")
    replace_run_text(data_row.cells[1].paragraphs[0], 0, "{{ item.equipment_name }}")
    replace_run_text(data_row.cells[2].paragraphs[0], 0, "{{ item.qty }}")
    replace_run_text(data_row.cells[3].paragraphs[0], 0, "{{ item.serial_number }}")
    replace_run_text(data_row.cells[4].paragraphs[0], 0, "{{ item.status }}")

    total_row = t3.rows[4]
    replace_run_text(total_row.cells[2].paragraphs[0], 0, "{{ total_qty }}")

    # --- Table 4: signatures --------------------------------------------------
    # Both signature names default to the corresponding party's name entered
    # above (receiving party / ICT rep) but are separate, independently
    # editable fields, in case the actual signer differs from the person on
    # the paperwork above.
    t4 = doc.tables[4]

    prepared_by_cell = t4.rows[2].cells[2]
    replace_run_text(prepared_by_cell.paragraphs[3], 0, "{{ signature_prepared_by }}")

    receiving_sig_cell = t4.rows[2].cells[0]
    receiving_sig_cell.add_paragraph()
    receiving_sig_cell.add_paragraph()
    name_p = receiving_sig_cell.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run("{{ signature_receiving_name }}")
    name_run.bold = True

    # --- Font pass: force Times New Roman everywhere ---------------------------
    # The source file sets Times New Roman explicitly on most runs, but relies
    # on the document's fallback theme font (Calibri) in a few places - the
    # equipment table body, the "WE CONFIRM THAT" paragraphs, and now our own
    # newly-added signature run - which would otherwise print in a visibly
    # different typeface. The one deliberate exception is the checkbox glyph
    # runs (MS Gothic), which need that font for the □/☑/☐ characters to
    # render at all; those are left untouched.
    def _iter_all_runs():
        for p in doc.paragraphs:
            yield from p.runs
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield from p.runs

    for run in _iter_all_runs():
        if run.font.name != "MS Gothic":
            run.font.name = "Times New Roman"

    doc.save(OUT_PATH)
    print("Wrote", os.path.abspath(OUT_PATH))


if __name__ == "__main__":
    main()
